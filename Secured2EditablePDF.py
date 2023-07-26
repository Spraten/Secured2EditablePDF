import os
import shutil
import subprocess
import sys
import re
import traceback
import json
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert
from colorama import Fore, init
from docx.shared import RGBColor
from alive_progress import alive_bar, alive_it

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        __import__(package)

packages = ['PyPDF2', 'pdf2docx', 'docx', 'docx2pdf', 'colorama', 'alive_progress']
for package in packages:
    install_and_import(package)

init(autoreset=True)

def decrypt_pdf(input_pdf_path, output_pdf_path, password):
    reader = PdfReader(input_pdf_path)
    if reader.is_encrypted:
        try:
            reader.decrypt(password)
            writer = PdfWriter()
            for page in alive_it(reader.pages):
                # print(f'writer.add_page({page})')
                writer.add_page(page)
            writer.write(output_pdf_path)
        except NotImplementedError:
            print(Fore.RED + f"Sorry, the encryption used in this PDF: {input_pdf_path} is not supported.")
    else:
        shutil.copy(input_pdf_path, output_pdf_path)

def convert_pdf_to_docx(input_pdf_path, output_docx_path):
    cv = Converter(input_pdf_path)
    try:
        cv.convert(output_docx_path, start=2, end=None)
    except ZeroDivisionError:
        pass
    cv.close()

def process_pdf_files(dir_path, password):
    pdf_files = [f for f in os.listdir(dir_path) if f.endswith('.pdf')]
    for pdf_file in pdf_files:
        print(f'Processing {pdf_file} pages:')
        input_pdf_path = os.path.join(dir_path, pdf_file)
        output_pdf_path = os.path.join(dir_path, 'decrypted', pdf_file[:-4] + '_decrypted.pdf')
        output_docx_path = os.path.join(dir_path, 'decrypted', pdf_file[:-4] + '.docx')
        if not os.path.exists(os.path.join(dir_path, 'decrypted')):
            os.makedirs(os.path.join(dir_path, 'decrypted'))
        decrypt_pdf(input_pdf_path, output_pdf_path, password)
        convert_pdf_to_docx(output_pdf_path, output_docx_path)

def docx_find_replace_text(doc_obj, find_text, replace_text):
    pattern = re.compile(find_text, re.IGNORECASE)
    for p in doc_obj.paragraphs:
        for run in p.runs:
            if pattern.match(run.text):
                run.text = re.sub(pattern, replace_text, run.text)

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_find_replace_text(cell, find_text, replace_text)

def docx_find_replace_white_text(doc_obj):
    # RGB value for white color
    white_color_rgb = (255, 255, 255)
    # RGB value for light gray color (adjust as needed)
    light_gray_color_rgb = (192, 192, 192)

    # Create RGBColor objects for white and light gray
    white_color = RGBColor(*white_color_rgb)
    light_gray_color = RGBColor(*light_gray_color_rgb)

    for p in doc_obj.paragraphs:
        for run in p.runs:
            font_color_rgb = run.font.color.rgb if run.font.color else None
            if font_color_rgb == white_color_rgb:
                run.font.color.rgb = light_gray_color

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        font_color_rgb = run.font.color.rgb if run.font.color else None
                        if font_color_rgb == white_color_rgb:
                            run.font.color.rgb = light_gray_color

def process_docx_files(dir_path, replacements):
    docx_files = [f for f in os.listdir(dir_path) if f.endswith('.docx')]
    for docx_file in docx_files:
        docx_path = os.path.join(dir_path, docx_file)
        try:
            doc = Document(docx_path)
            for replacement in replacements:
                find_text = replacement['find_text']
                replace_text = replacement['replace_text']
                docx_find_replace_white_text(doc)  # Call to replace white text with light gray
                docx_find_replace_text(doc, find_text, replace_text)
            doc.save(docx_path)
            # docx_find_replace_text(doc, find_text, replace_text)
            new_docx_filename = re.sub(r"_\d+", "", docx_file.replace("_decrypted", ""))
            modified_docx_path = os.path.join(dir_path, new_docx_filename)
            doc.save(modified_docx_path)
            new_pdf_filename = new_docx_filename.replace(".docx", ".pdf")
            modified_pdf_path = os.path.join(dir_path, new_pdf_filename)
            convert(modified_docx_path, modified_pdf_path)
            course_name = new_docx_filename.split(' - ')[0]
            target_docx_dir = os.path.join(dir_path, '..', f'{course_name} Docx')
            target_pdf_dir = os.path.join(dir_path, '..', f'{course_name} PDF')
            if not os.path.exists(target_docx_dir):
                os.makedirs(target_docx_dir)
            if not os.path.exists(target_pdf_dir):
                os.makedirs(target_pdf_dir)
            try:
                shutil.move(modified_docx_path, os.path.join(target_docx_dir, new_docx_filename))
                shutil.move(modified_pdf_path, os.path.join(target_pdf_dir, new_pdf_filename))
            except PermissionError as e:
                print(f"{Fore.RED}An error occurred while moving the files:")
                print(f"{e}")
                print(f"{Fore.YELLOW}Please check if the file is open and close it before re-running the script.")
                retry = input(f"{Fore.YELLOW}Enter 'y' to retry moving the file once it is closed: ")
                if retry.lower() == 'y':
                    process_docx_files(dir_path, replacements)
        except PermissionError as e:
            print(f"{Fore.RED}An error occurred while processing the file:")
            print(f"{e}")

def read_find_replace_from_json(json_path):
    try:
        with open(json_path, 'r') as f:
            replacements = json.load(f)
    except FileNotFoundError:
        print(f"JSON file at path {json_path} was not found. Creating a new one.")
        replacements = [{'find_text': '', 'replace_text': ''}]
        with open(json_path, 'w') as f:
            json.dump(replacements, f)
    else:
        return replacements
        
def main():
    dir_path = input(Fore.GREEN + "Please enter the absolute path of the PDFs (leave blank for current directory): ")
    if not dir_path:
        dir_path = '.'

    if not os.path.exists(dir_path):
        print(Fore.RED + f"The directory {dir_path} does not exist.")
        return

    pdf_files = [f for f in os.listdir(dir_path) if f.endswith('.pdf')]

    if not pdf_files:
        print(Fore.RED + f"No PDF files found in the directory: {dir_path}")
        return

    password = input(Fore.GREEN + "Please enter the password (leave blank if password is stored in password.txt): ")
    if not password:
        with open('password.txt', 'r') as f:
            password = f.read().strip()
    try:
        process_pdf_files(dir_path, password)
    except Exception as e:
        print(Fore.RED + f"An unexpected error occurred while processing PDF files: {e}")
        traceback.print_exc()
        return
    while True:
        if input(Fore.GREEN + "Do you have a JSON file for find and replace pairs? (Y/n): ")[:1].lower().strip() != 'n':
            find_replace_json = input(Fore.GREEN + "Enter the location of the JSON file containing find/replace strings (default: find_replace.json): ") or "find_replace.json"
            if not os.path.isfile(find_replace_json):
                print(Fore.RED + f"The file {find_replace_json} does not exist.")
                continue
            replacements = read_find_replace_from_json(find_replace_json)
            break
        else:
            find_text = input(Fore.GREEN + "Please enter the text to find (default: (Licensed To:).*") or r"(Licensed To:).*"
            replace_text = input(Fore.GREEN + "Please enter the text to replace it with (default: ===Wow_Something_used_to_be_here===): ") or "===Wow_Something_used_to_be_here==="
            replacements = [{'find_text': find_text, 'replace_text': replace_text}]
            break
    try:
        process_docx_files(os.path.join(dir_path, 'decrypted'), replacements)
    except Exception as e:
        print(Fore.RED + f"An unexpected error occurred while processing DOCX files: {e}")
        traceback.print_exc()
        return

if __name__ == "__main__":
    main()
